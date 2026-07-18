from pathlib import Path

import docx
from langchain_community.document_loaders import DirectoryLoader, PyPDFLoader, TextLoader
from langchain_core.documents import Document

BASE_DIR = Path(__file__).resolve().parents[3]
KNOWLEDGE_BASE = BASE_DIR / "knowledge_base"


def load_markdown():
    loader = DirectoryLoader(
        str(KNOWLEDGE_BASE),
        glob="**/*.md",
        loader_cls=TextLoader,
    )
    return loader.load()


def load_txt():
    loader = DirectoryLoader(
        str(KNOWLEDGE_BASE),
        glob="**/*.txt",
        loader_cls=TextLoader,
    )
    return loader.load()


def load_pdf():
    loader = DirectoryLoader(
        str(KNOWLEDGE_BASE),
        glob="**/*.pdf",
        loader_cls=PyPDFLoader,
    )
    return loader.load()


def load_docx():
    # Deliberately not UnstructuredWordDocumentLoader: it needs the heavy
    # `unstructured[docx]` package (not installed - see requirements.txt
    # history). python-docx gives plain paragraph text, which is all
    # chunking/embedding needs, for a much lighter dependency.
    documents = []

    for path in KNOWLEDGE_BASE.rglob("*.docx"):
        try:
            doc = docx.Document(str(path))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception:
            continue

        if text.strip():
            documents.append(Document(page_content=text, metadata={"source": str(path)}))

    return documents


def load_all_documents():
    documents = []

    documents.extend(load_markdown())

    try:
        documents.extend(load_txt())
    except Exception:
        pass

    try:
        documents.extend(load_pdf())
    except Exception:
        pass

    try:
        documents.extend(load_docx())
    except Exception:
        pass

    return documents
